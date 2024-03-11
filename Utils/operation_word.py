#-- coding: utf-8 --

#@Time : 2022/8/22 10:39

#@Author : tianxh

#@Email : tianxh@casking.com.cn

#@File : operation_word.py

#@Software: PyCharm

import docx
import sys,os
from Utils.operationyaml import *
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.shared import RGBColor
import pypandoc
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
from Utils.cennect_mysql_zentao import *
from Utils.operation_zentao_mysql import *
from Utils.currenttime import *
import ddddocr

sys.path.append((os.path.abspath(os.path.join(os.path.dirname(__file__), '../'))))
# html文档的位置
# html_path = sys.path[1]+"/2022_07_07_16_44_55result.html"
# # 转换生成word文档的位置
# word_path = sys.path[1]+"/ceshi_to_word.docx"
# print(html_path)
# print(word_path)
# pypandoc.convert_file(html_path, 'docx', outputfile=word_path)
# ceshi_word = sys.path[1]+"/XXX测试报告模板.docx"
# ceshi_phone = sys.path[1]+"/下载.jpg"
# doc = docx.Document(ceshi_word)
# doc_ceshi = docx.Document(word_path)
# print(len(doc.paragraphs))
# print(doc.paragraphs[4].text)

def html_to_word():
    """
    html文件转word
    :return:
    """
    html_path = sys.path[1]+"/-.html"
    # # 转换生成word文档的位置
    word_path = sys.path[1] + "/ceshi_to_word.docx"
    pypandoc.convert_file(html_path, 'docx', outputfile=word_path)

def add_head(word_file):
    """
    增加word标题
    :param word_file: word文件
    :return:
    """
    doc = docx.Document(word_file)
    doc.add_heading('身份认证系统接口自动化测试报告', 2)
    doc.add_paragraph('身份认证系统接口自动化测试报告详情:')
    doc.save(word_file)



def merge_word(files,final_docx):
    """
    word文件合并
    :param files: 文件名列表
    :param final_docx: 合并后的文件路径
    :return:
    """
    new_document = Document()
    composer = Composer(new_document)
    for fn in files:
        composer.append(Document(fn))
        composer.save(final_docx)