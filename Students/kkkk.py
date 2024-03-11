#-- coding: utf-8 --

#@Time : 2022/7/1 10:06

#@Author : tianxh

#@Email : tianxh@casking.com.cn

#@File : kkk.py

#@Software: PyCharm

import pypandoc
import wget
# -*- coding:utf-8 -*-
import pypandoc
import docx
import sys,os
import yagmail
from Utils.operationyaml import *
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.shared import RGBColor

from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
from Utils.cennect_mysql_zentao import *
from Utils.operation_zentao_mysql import *
from Utils.currenttime import *
import ddddocr

sys.path.append((os.path.abspath(os.path.join(os.path.dirname(__file__), '../'))))
# html文档的位置
# html_path = sys.path[1]+"/2022_07_07_16_44_55result.html"
# # 转换生成word文档的位置
word_path = sys.path[1]+"/ceshi_to_word.docx"
# print(html_path)
# print(word_path)
# pypandoc.convert_file(html_path, 'docx', outputfile=word_path)
ceshi_word = sys.path[1]+"/XXX测试报告模板.docx"
ceshi_phone = sys.path[1]+"/下载.jpg"
# doc = docx.Document(ceshi_word)
# doc_ceshi = docx.Document(word_path)
# print(len(doc.paragraphs))
# print(doc.paragraphs[4].text)
#
# doc.add_heading('身份认证系统接口自动化测试报告',2)
# # 换行
# doc.paragraphs[0].runs[0].add_break()
# doc.add_picture(ceshi_phone,
# width=docx.shared.Inches(3),
# height=docx.shared.Cm(4))
#
# p = doc.add_paragraph('A plain paragraph having some')#在段落后面追加文本，并可设置样式
# p.add_run('bold').bold =True
# p.add_run('and some')
# p.add_run('italic.').italic =True

# doc.add_heading('Heading, level 1', level=1)
# doc.add_paragraph('Intense quote', style='Intense Quote')#添加项目列表（前面一个小圆点）
# doc.add_paragraph('first item in unordered list', style='List Bullet')
# doc.add_paragraph('second item in unordered list', style='List Bullet')#添加项目列表（前面数字）
# doc.add_paragraph('first item in ordered list', style='List Number')
# doc.add_paragraph('second item in ordered list', style='List Number')#添加图片

# records=(
# (3, '101', 'Spam'),
# (7, '422', 'Eggs'),
# (4, '631', 'Spam, spam, eggs, and spam')
# )
# #添加表格：一行三列#表格样式参数可选：#Normal Table#Table Grid#Light Shading、 Light Shading Accent 1 至 Light Shading Accent 6#Light List、Light List Accent 1 至 Light List Accent 6#Light Grid、Light Grid Accent 1 至 Light Grid Accent 6#太多了其它省略...
# table = doc.add_table(rows=1, cols=3)#获取第一行的单元格列表
# hdr_cells =table.rows[0].cells#下面三行设置上面第一行的三个单元格的文本值
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# for qty, id, desc in records:#表格添加行，并返回行所在的单元格列表
#     row_cells =table.add_row().cells
#     row_cells[0].text=str(qty)
#     row_cells[1].text =id


# p1 = doc.add_paragraph('新页新段落')
#
# # 添加一个 2×2 表格
# table = doc.add_table(rows=2, cols=2)
# # 获取第1行第2列单元格
# cell = table.cell(0, 1)
#
# # 设置单元格文本
# cell.text = '第1行第2列'
#
# # 获取第2行
# row = table.rows[1]
# row.cells[0].text = '橡皮擦'
# row.cells[1].text = '乔喻'


# 一级标题 level=0
# head0 = doc.add_heading(level=0)
# # 标题居中
# head0.alignment = WD_ALIGN_PARAGRAPH.CENTER
# title_run = head0.add_run('这是一个居中的标题', )
# title_run.font.size = Pt(24)
# 标题英文字体
# title_run = head0.add_run('这是一个居中的标题', )
# title_run.font.size = Pt(24)
# title_run.font.name = 'Times New Roman'
# 标题中文字体
# title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
# # 字体颜色
# title_run.font.color.rgb = RGBColor(4, 60, 169)


# style = doc.styles.add_style('textstylew', WD_STYLE_TYPE.PARAGRAPH)
# # 字体大小
# style.font.size = Pt(20)
# # 字体颜色
# style.font.color.rgb = RGBColor(66, 100, 0)
# p1 = doc.add_paragraph('神人橡皮擦',style=style)


# for i,para in enumerate(doc.paragraphs):
#     print(f'No.{i+1}', para.text, sep='')

# with open(word_path, 'rb', encoding="utf-8") as f:
#     # yaml文件中读取内容
#     print(f.read())

# for p in doc_ceshi.paragraphs:
#     print(p.text)


# p1 = doc.add_paragraph('这是一个段落')
# p1.add_run('加粗的一句话').bold = True
# doc.add_paragraph('这是第二个段落')
# doc.save(ceshi_word)

# print(len(doc_ceshi.paragraphs))
# print(doc_ceshi.paragraphs[0].text)
#
#
#
#
# tb = doc.tables[0].rows[0].cells[2].text='是多少'
# # doc.save(ceshi_word)
# print(tb)


# from Utils.operation_word import *
# html_to_word()
# ceshi_word = sys.path[1]+"/XXX测试报告模板.docx"
# test_report_word = sys.path[1]+"/测试报告.docx"
# # add_head(ceshi_word)
# merge_word([ceshi_word,word_path],test_report_word)


# import re
# from docx import Document
# from docx.shared import Pt
# from docx.shared import Inches
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
#
# alignment_dict={'justify':WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
#                 'center':WD_PARAGRAPH_ALIGNMENT.CENTER,
#                 'right':WD_PARAGRAPH_ALIGNMENT.RIGHT,
#                 'left':WD_PARAGRAPH_ALIGNMENT.LEFT}
# def get_docx_align(doc,align='justify'):
#     for para in doc.paragraphs:
#         inline=para.runs
#         for i in range(len(inline)):
#             inline[i].text=re.sub('[\t]+','',inline[i].text)
#             inline[i].text=re.sub('[\n]+','\n',inline[i].text)
#             inline[i].text=re.sub('[" "]+',' ',inline[i].text)
#             #p=para._element
#             #p.getparent().remove(p)
#             #p._p=p._element=None
#         paragraph=para
#         paragraph_format=paragraph.paragraph_format
#         paragraph_format.alignment=alignment_dict.get(align.lower())
# doc=Document(word_path)
# get_docx_align(doc,align='justify')
# doc.save(word_path)

# import pdfkit
# html_path = sys.path[1]+"/-.html"
# pdfkit.from_file(html_path, sys.path[1]+"/result.pdf")

import xhtml2pdf
# def create_pdf(pdf_data):
#     filename = app.config['UPLOAD_FOLDER'] + "/file.pdf"
#     f = file(filename, "wb")
#     pdf = pisa.CreatePDF(StringIO(pdf_data.encode('utf-8')), f)
#     return pdf
#
# pdf = create_pdf(render_template('receipt.htm', entry=entry, filename=filename))
#
#
# def fetch_resources(uri, rel):
#     if uri.startswith(settings.MEDIA_URL):
#         path = os.path.join(settings.MEDIA_ROOT,uri.replace(settings.MEDIA_URL, ""))
#     elif uri.startswith(settings.STATIC_URL):
#         path = os.path.join(settings.LOADING_STATIC_FOR_PDF,uri.replace(settings.STATIC_URL, ""))
#     return path
#
# def preview_pdf(request):
#     context = RequestContext(request)
#     if request.method == 'GET':
#         context = {
#             # sending data to html
#         }
#         rendered_html = render_to_string("path/xxx.html", locals())
#         template = get_template('path/xxx.html')
#         html = template.render(context)
#         result = StringIO.StringIO()
#         file = open('test.pdf', "w+b")
#         pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=file, encoding='utf-8',link_callback=fetch_resources)
#         file.seek(0)
#         pdf = file.read()
#         file.close()
#         return HttpResponse(pdf, 'application/pdf')

from xhtml2pdf import pisa
import urllib3
import urllib.request

# url=urllib.request.urlopen(sys.path[1]+"/-.html")
# print("http://"+sys.path[1]+"/-.html")

# sourceHtml=url.read()
# pisa.showLogging()
#
# outputFilename = sys.path[1]+"/test.pdf"
#
# def convertHtmlToPdf(sourceHtml, outputFilename):
#     resultFile = open(outputFilename, "w+b")
#     pisaStatus = pisa.CreatePDF(sourceHtml,resultFile)
#     resultFile.close()
#     return pisaStatus.err
#
# if __name__=="__main__":
#     pisa.showLogging()
#     convertHtmlToPdf(sourceHtml, outputFilename)

# def convert_html_to_pdf(source_html, output_filename):
#
#     result_file = open(output_filename, "w+b")
#     print(result_file)
#     # convert HTML to PDF
#     pisa_status = pisa.CreatePDF(
#             source_html,                # the HTML to convert
#             dest=result_file)           # file handle to recieve result
#
#     # close output file
#     result_file.close()                 # close output file
#
# # return True on success and False on errors
#     return pisa_status.err
# if __name__=="__main__":
#     convert_html_to_pdf(sys.path[1]+"/-.html",outputFilename)

# import requests
# import re
# import hashlib
# from requests_toolbelt import MultipartEncoder
#
# refre_url = 'http://192.168.13.148:63802/user-refreshRandom.html'
# r = requests.get(refre_url)
# verify_rand = str(json.loads(json.dumps(r.text)))
# s = requests.session()
# user = 'tianxh'
# password = 'Casking@1234'
# # password = '123456'
# url = 'http://192.168.13.148:63802/user-login.html'
# r = s.get(url)
# # print(r.content.decode('utf-8'))
#
# # print(r.content.decode('utf-8'))
# verify = re.findall(r"name='referer' id='referer' value='(.*?)'  />", r.content.decode('utf-8'))[0]
# print(verify)
# # 第一次加密密码
# pwd1md5 = hashlib.md5()
# pwd1md5.update(password.encode('utf-8'))
# pwd1_result = pwd1md5.hexdigest()
# print(pwd1_result)
# print(pwd1_result+verify_rand)
# # 第2次加密
# pwd2md5 = hashlib.md5()
# pwd2md5.update((pwd1_result+verify_rand).encode('utf-8'))
# pwd2_result = pwd2md5.hexdigest()
# print(pwd2_result)
# print(int(verify_rand))
# body = {
#                  "account": user,
#                "password": pwd2_result,
#                "passwordStrength": 2,
#                "referer": "/",
#                "verifyRand": verify_rand,
#                "keepLogin": 1,
#                 "captcha" : ""
#
#               #  "account":user,
#               # "password":'c450fff1edd5346abca25254be1e2d4c',
#               # "keepLogin[]":"on",
#               #   "referer":"http://127.0.0.1/zentao/my/"
#                }
# bodys = 'account=tianxh&password='+pwd2_result+'&passwordStrength=2&referer=%252F&verifyRand='+verify_rand+'&keepLogin=1&captcha='
# print(body)
# headers = {
#     "User-Agent": "Mozilla/5.0 (Windows NT 10.0; WOW64; rv:44.0) Gecko/20100101 Firefox/44.0",
#     "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
#     "Accept-Language": "zh-CN,zh;q=0.8,en-US;q=0.5,en;q=0.3",
#     "Accept-Encoding": "gzip, deflate",
#     # "Cookie":  # 头部没登录前不用传cookie，因为这里cookie就是保持登录的
#     "Connection": "keep-alive",
#     "Content-Type": "application/x-www-form-urlencoded",
#     "Referer": "http://192.168.13.148:63802/my/"
#
#     }
#
# r = s.post('http://192.168.13.148:63802/user-login.html',headers, body)
# print(s)
# print(r.content.decode('utf-8'))
# # 访问测试页面
# headers = {
#     "User-Agent": "Mozilla/5.0 (Windows NT 10.0; WOW64; rv:44.0) Gecko/20100101 Firefox/44.0",
#     "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
#     "Accept-Language": "zh-CN,zh;q=0.8,en-US;q=0.5,en;q=0.3",
#     "Accept-Encoding": "gzip, deflate",
#     "Cookie":  "lang=zh-cn; device=desktop; theme=default; keepLogin=on; preBranch=0; preProductID=1; lastBugModule=1; goback=%7B%22qa%22%3A%22http%3A%5C%2F%5C%2F192.168.13.148%3A63802%5C%2Fbug-browse-1.html%22%7D; zentaosid=a5rfv0k2n1l9jsu3fqpg09cugp; bugModule=0; bugBranch=0; treeBranch=0; qaBugOrder=id_desc; za=tianxh; zp=e915b0ab58b574cac2210c93c58de8863026d494; tab=qa; windowWidth=693; windowHeight=544",
#     "Connection": "keep-alive",
#     "Content-Type": "application/x-www-form-urlencoded",
#     }
# test = requests.get("http://192.168.13.148:63802/my/",headers)
# print(test.content.decode('utf-8'))
# if "我的任务" in test.text:
#     print('登录成功！！')
# else:
#     print('登录失败！！')
# # 提交bug接口
# url2 = 'http://192.168.13.148:63802/pro/bug-create-1-0-moduleID=0.html'
# body2 = MultipartEncoder(
#     fields=[
#         ('product', "1"),
#         ('module', '0'),
#         ('project', ' '),
#         ('openedBuild[]', 'trunk'),
#         ('assignedTo', 'xiaojc'),
#         ('deadline', ''),
#         ('type', 'codeerror'),
#         ('os', ''),
#         ('browser', ''),
#         ('title', '正确的账号密码登录失败'),  # bug 名称
#         ('color', ''),
#         ('severity', '3'),
#         ('pri', '3'),
#         ('steps', '<p>[步骤]</p>\n<p>输入正确的账号名密码进行完成登录</p>\n<br />\n<p>[结果]</p>\n登录失败<br />\n<p>[期望]</p>\n登录成功<br />'),
#         ('story', '0'),
#         ('task','0'),
#         ('oldTaskID', '0'),
#         ('mailto[]', ''),
#         ('contactListMenu', ''),
#         ('keywords', ''),
#         ('status', 'active'),
#         ('labels[]', ''),
#         ('files[]', ''),
#         ('uid', '630d7242b50d7'),
#         ('case', '0'),
#         ('caseVersion', '0'),
#         ('caseVersion', '0'),
#         ('result', '0'),
#         ('testtask', '0'),
#             ],
#     )
# # 请求提交bug接口
# headers={'Content-Type': body2.content_type,
# "Cookie":"lang=zh-cn; device=desktop; theme=default; keepLogin=on; preBranch=0; preProductID=1; zentaosid=a5rfv0k2n1l9jsu3fqpg09cugp; bugBranch=0; treeBranch=0; qaBugOrder=id_desc; lastBugModule=2; za=tianxh; zp=450f378e65fa7beb1f658036c41e6aaabbc9c501; tab=qa; bugModule=0; windowWidth=748; windowHeight=544; goback=%7B%22qa%22%3A%22http%3A%5C%2F%5C%2F192.168.13.148%3A63802%5C%2Fbug-browse-1.html%22%7D"
#          }
# r2 = requests.post(url2, headers, body2)
# if '保存成功' in r2.text:
#     print('bug提交成功！')
# else:
#     print('bug提交失败')


# import pdfkit
# import wkhtmltopdf
# import os
# print(sys.path)
# filename = sys.path[1] + "/" + 'result.html'
# pdfkit.from_file(filename, 'C:\website2.pdf')

# os.system('wkhtmltopdf C:\pyproject\\reliangApiTest\\result.html C:\website1.pdf')

# import yagmail
# test_report = sys.path[0]+'/'+'test.zip'
# print(test_report)
# yagindex = yagmail.SMTP(user='cdzdhcs@casking.com.cn', password='Chandao123', host='smtp.exmail.qq.com')
# yagindex.send('tianxh@casking.com.cn', '身份认证系统性能测试报告', 'wwee', test_report)

# import xml.etree.ElementTree as ET
# def change_one_xml(xml_path, xml_dw, update_content):
#     # 打开xml文档
#     doc = ET.parse(xml_path)
#     root = doc.getroot()
#     # 查找修改路劲
#     sub1 = root.find(xml_dw)
#     sub1.text = update_content
#     # 修改标签内容
#     sub1.text = update_content
#     # 保存修改
#     doc.write(xml_path)
#
# if __name__ == '__main__':
#     print(sys.path)
#     xml_path = sys.path[1]+'/'+'p_udemr'+'/'+'test_cases'+'/'+'电子病历调阅系统.jmx'
#
#     # 修改文件中的xpath定位
#     # xml_dw = './/country[@name="Singapore"]/year'
#     xml_dw = './/stringProp[@name="Header.name"][text()="Authorization"]'
#     # 想要修改成什么内容
#     update_content = '9999'
#     change_one_xml(xml_path, xml_dw, update_content)

# from xml.etree import ElementTree as ET
# # 1、打开xml文件
# xml_path = sys.path[1]+'/'+'p_udemr'+'/'+'test_cases'+'/'+'电子病历调阅系统.jmx'
# print(xml_path)
# tree =ET.parse(xml_path)
# # 获xml文件的内容取根标签
# root = tree.getroot()
# print(root)
# xml_dw = './/stringProp[text()="Authorization"]'
# print(root.xpath(xml_dw))

# import lxml.etree
# xml_path = sys.path[1]+'/'+'p_udemr'+'/'+'test_cases'+'/'+'电子病历调阅系统.jmx'
# root = lxml.etree.parse(xml_path)
#
# xml_dw = './/stringProp[contains(text(),"Bearer ")]'
# # xml_dw = './/stringProp/text()="Bearer c8e9572e-fdfe-45b5-92af-ea52243d0de6"'
#
# for i in range(0,len(root.xpath(xml_dw))):
#     root.xpath(xml_dw)[i].text = 'Bearer jjjjjjjjjjjjj'
#     print(root.xpath(xml_dw)[i].text)
#
#
# root.write(xml_path)


import json
#
# with open(jmeter_script_config) as f:
#     superHeroSquad = json.load(f)
#
# print(type(superHeroSquad))  # Output: dict
# print(superHeroSquad.keys())

# import pandas as pd
# df = pd.read_json(jmeter_script_config)
# print(df)
# from Utils.operation_remote_server import *
# # comd = 'appscancmd'
# comd = 'net time'
# execute_linux_commnd('192.168.8.61' ,'Administrator', 'lw@2019', comd)

# import winrm
# wintest = winrm.Session('192.168.8.61',auth=('Administrator','lw@2019'))
# ret = wintest.run_cmd('ipconfig')
# print(ret.std_out)

# bug_to_zentao(False, 'ffff', 'sds', 'zhaop', 'xiaojc')
import requests
response = requests.request(
                method='get',
                url='http://baike.baidu.com/api/openapi/BaikeLemmaCardApi?scope=103&format=json&appid=379020&bk_key=关键字&bk_length=600用例',
                data={},
                headers={}
            )
response.encoding='utf-8'
print(response.text)